import os
import re
import time
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


_EPISODE_RE = re.compile(r"^第\d+集[：:]")
_SCENE_RE = re.compile(r"^(?:场景号[：:]\s*\d+|【?场景\s*\d+)")
_SEPARATOR_RE = re.compile(r"^={5,}$")
_META_LABEL_RE = re.compile(
    r"^(角色造型&情绪|角色造型|动作描述|对白|分镜提示|内外景|时间|地点|场景号)[：:]"
)


class ExportService:
    def __init__(self, output_dir: str) -> None:
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)
        self.pdf_font_name = self._register_fonts()

    def _register_fonts(self) -> str:
        try:
            pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
            return "STSong-Light"
        except Exception:
            pass
        possible_fonts = [
            r"C:\Windows\Fonts\msyh.ttf",
            r"C:\Windows\Fonts\simhei.ttf",
            r"C:\Windows\Fonts\simsun.ttc",
        ]
        for font_path in possible_fonts:
            if os.path.exists(font_path):
                try:
                    pdfmetrics.registerFont(TTFont("ChineseFont", font_path))
                    return "ChineseFont"
                except Exception:
                    continue
        return "Helvetica"

    # ---- DOCX ----

    def export_docx(self, title: str, content: str) -> str:
        path = os.path.join(self.output_dir, f"{self._safe_name(title)}_{int(time.time())}.docx")
        normalized_title = self._normalize_text(title)
        normalized_content = self._normalize_text(content)
        document = Document()

        heading = document.add_heading(normalized_title, level=1)
        self._set_docx_font(heading, size=18)

        for line in normalized_content.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if _SEPARATOR_RE.match(stripped):
                document.add_page_break()
                continue
            if _EPISODE_RE.match(stripped):
                h = document.add_heading(stripped, level=2)
                self._set_docx_font(h, size=14)
                continue
            if _SCENE_RE.match(stripped):
                h = document.add_heading(stripped, level=3)
                self._set_docx_font(h, size=12)
                continue
            paragraph = document.add_paragraph()
            meta_match = _META_LABEL_RE.match(stripped)
            if meta_match:
                label = meta_match.group(0)
                rest = stripped[len(label):]
                run_bold = paragraph.add_run(label)
                run_bold.bold = True
                self._set_run_font(run_bold, size=11)
                if rest:
                    run_normal = paragraph.add_run(rest)
                    self._set_run_font(run_normal, size=11)
            else:
                run = paragraph.add_run(stripped)
                self._set_run_font(run, size=11)

        document.save(path)
        return path

    def _set_docx_font(self, element, size: int = 11) -> None:
        for run in element.runs:
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(size)

    def _set_run_font(self, run, size: int = 11) -> None:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(size)

    # ---- PDF ----

    def export_pdf(self, title: str, content: str) -> str:
        path = os.path.join(self.output_dir, f"{self._safe_name(title)}_{int(time.time())}.pdf")
        normalized_title = self._normalize_text(title)
        normalized_content = self._normalize_text(content)

        pdf = canvas.Canvas(path, pagesize=A4)
        font_name = self.pdf_font_name
        page_width, page_height = A4
        left_margin = 40
        right_margin = 40
        top_margin = 48
        bottom_margin = 48
        usable_width = page_width - left_margin - right_margin

        try:
            pdf.setFont(font_name, 11)
        except Exception:
            font_name = "Helvetica"
            pdf.setFont(font_name, 11)

        y = page_height - top_margin

        def _new_page():
            nonlocal y
            pdf.showPage()
            pdf.setFont(font_name, 11)
            y = page_height - top_margin

        def _draw_line(text: str, font_size: int = 11, line_height: int = 16, bold_prefix: str = ""):
            nonlocal y
            if y < bottom_margin + line_height:
                _new_page()
            pdf.setFont(font_name, font_size)
            wrapped = self._wrap_pdf_line(text, font_name, font_size, usable_width)
            if not wrapped:
                wrapped = [""]
            for wl in wrapped:
                if y < bottom_margin + line_height:
                    _new_page()
                    pdf.setFont(font_name, font_size)
                pdf.drawString(left_margin, y, wl)
                y -= line_height

        # Title
        _draw_line(normalized_title, font_size=16, line_height=24)
        y -= 8

        for line in normalized_content.splitlines():
            stripped = line.strip()
            if not stripped:
                y -= 6
                continue
            if _SEPARATOR_RE.match(stripped):
                _new_page()
                continue
            if _EPISODE_RE.match(stripped):
                y -= 10
                _draw_line(stripped, font_size=14, line_height=22)
                y -= 4
                continue
            if _SCENE_RE.match(stripped):
                y -= 6
                _draw_line(stripped, font_size=12, line_height=18)
                continue
            _draw_line(stripped, font_size=11, line_height=16)

        pdf.save()
        return path

    # ---- Helpers ----

    def _normalize_text(self, value: str) -> str:
        return value.replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n")

    def _wrap_pdf_line(self, text: str, font_name: str, font_size: int, max_width: float) -> list[str]:
        if not text:
            return [""]
        lines: list[str] = []
        current = ""
        for char in text:
            candidate = f"{current}{char}"
            width = pdfmetrics.stringWidth(candidate, font_name, font_size)
            if width <= max_width:
                current = candidate
            else:
                if current:
                    lines.append(current)
                current = char
        if current:
            lines.append(current)
        return lines

    def _safe_name(self, value: str) -> str:
        cleaned = value.replace("/", "_").replace("\\", "_").replace(":", "_").replace("*", "_").replace("?", "_").replace('"', "_").replace("<", "_").replace(">", "_").replace("|", "_")
        return cleaned[:30] or "script"
